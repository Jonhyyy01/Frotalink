from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("entregas_pap/Relatorio_Final_Frotalink_PAP.docx")


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    for idx, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        set_cell_text(table.rows[0].cells[idx], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=10)
    doc.add_paragraph()
    return table


def add_para(doc, text="", style=None, align=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True, color="555555")
    return p


def add_figure_placeholder(doc, title, note="Inserir captura de ecra ou imagem correspondente."):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(15.5)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(note)
    set_run_font(r, size=10.5, italic=True, color="666666")
    add_caption(doc, title)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(code)
    set_run_font(r, name="Consolas", size=9, color="1F4D78")


def page_break(doc):
    doc.add_page_break()


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 16)
        st.paragraph_format.space_after = Pt(6)


def apply_portuguese_accents(doc):
    replacements = {
        "Relatorio": "Relatório",
        "relatorio": "relatório",
        "Relatorios": "Relatórios",
        "relatorios": "relatórios",
        "Indice": "Índice",
        "indice": "índice",
        "Aptidao": "Aptidão",
        "aptidao": "aptidão",
        "Informacao": "Informação",
        "informacao": "informação",
        "Configuracao": "Configuração",
        "configuracao": "configuração",
        "Descricao": "Descrição",
        "descricao": "descrição",
        "Obtencao": "Obtenção",
        "obtencao": "obtenção",
        "Qualificacao": "Qualificação",
        "qualificacao": "qualificação",
        "Nivel": "Nível",
        "nivel": "nível",
        "Informatica": "Informática",
        "informatica": "informática",
        "DECLARACAO": "DECLARAÇÃO",
        "DECLARACAO DE COMPROMISSO": "DECLARAÇÃO DE COMPROMISSO",
        "E autorizada": "É autorizada",
        "e original": "é original",
        "a OFICINA": "à OFICINA",
        "as principais": "às principais",
        "as funcionalidades": "às funcionalidades",
        "a sua funcao": "à sua função",
        "a base de dados": "à base de dados",
        "as varias": "às várias",
        "as areas": "às áreas",
        "a escola": "à escola",
        "a todos": "a todos",
        "a gestao": "à gestão",
        "as operacoes": "às operações",
        "unico": "único",
        "Unico": "Único",
        "util": "útil",
        "Util": "Útil",
        "informacoes": "informações",
        "Informacoes": "Informações",
        "operacao": "operação",
        "Operacao": "Operação",
        "revisoes": "revisões",
        "Revisoes": "Revisões",
        "comunicacao": "comunicação",
        "Comunicacao": "Comunicação",
        "atraves": "através",
        "Atraves": "Através",
        "execucao": "execução",
        "Execucao": "Execução",
        "sessao": "sessão",
        "Sessao": "Sessão",
        "Gestao": "Gestão",
        "gestao": "gestão",
        "Operacao": "Operação",
        "operacao": "operação",
        "Logistica": "Logística",
        "logistica": "logística",
        "Introducao": "Introdução",
        "introducao": "introdução",
        "Proposito": "Propósito",
        "proposito": "propósito",
        "Ambito": "Âmbito",
        "ambito": "âmbito",
        "Definicoes": "Definições",
        "definicoes": "definições",
        "Acronimos": "Acrónimos",
        "acronimos": "acrónimos",
        "Abreviaturas": "Abreviaturas",
        "Visao": "Visão",
        "visao": "visão",
        "Identificacao": "Identificação",
        "identificacao": "identificação",
        "Publico": "Público",
        "publico": "público",
        "Necessarios": "Necessários",
        "necessarios": "necessários",
        "Planeamento": "Planeamento",
        "Criacao": "Criação",
        "criacao": "criação",
        "Edicao": "Edição",
        "edicao": "edição",
        "Eliminacao": "Eliminação",
        "eliminacao": "eliminação",
        "Utilizacao": "Utilização",
        "utilizacao": "utilização",
        "Aplicacao": "Aplicação",
        "aplicacao": "aplicação",
        "Autenticacao": "Autenticação",
        "autenticacao": "autenticação",
        "Permissoes": "Permissões",
        "permissoes": "permissões",
        "Protecao": "Proteção",
        "protecao": "proteção",
        "Sessoes": "Sessões",
        "sessoes": "sessões",
        "Ligacao": "Ligação",
        "ligacao": "ligação",
        "Base de Dados": "Base de Dados",
        "Manutencoes": "Manutenções",
        "manutencoes": "manutenções",
        "Avarias": "Avarias",
        "Abastecimentos": "Abastecimentos",
        "Combustivel": "Combustível",
        "combustivel": "combustível",
        "Veiculos": "Veículos",
        "veiculos": "veículos",
        "Viaturas": "Viaturas",
        "Motoristas": "Motoristas",
        "Condutores": "Condutores",
        "Codigo": "Código",
        "codigo": "código",
        "Paginas": "Páginas",
        "paginas": "páginas",
        "Pagina": "Página",
        "pagina": "página",
        "ecra": "ecrã",
        "Ecrã": "Ecrã",
        "Formularios": "Formulários",
        "formularios": "formulários",
        "Validacao": "Validação",
        "validacao": "validação",
        "Tecnica": "Técnica",
        "tecnica": "técnica",
        "Documentacao": "Documentação",
        "documentacao": "documentação",
        "Tecnico": "Técnico",
        "tecnico": "técnico",
        "Automatica": "Automática",
        "automatica": "automática",
        "Critica": "Crítica",
        "critica": "crítica",
        "Apreciacao": "Apreciação",
        "Conclusao": "Conclusão",
        "conclusao": "conclusão",
        "Bibliografia": "Bibliografia",
        "Endereco eletronico": "Endereço eletrónico",
        "reproducao": "reprodução",
        "declaracao": "declaração",
        "compromete": "compromete",
        "Antiplagio": "Antiplágio",
        "ANTIPLAGIO": "ANTIPLÁGIO",
        "citacoes": "citações",
        "contribuicoes": "contribuições",
        "estao": "estão",
        "consciencia": "consciência",
        "utilizacao": "utilização",
        "etica": "ética",
        "desclassificacao": "desclassificação",
        "realizacao": "realização",
        "orientacao": "orientação",
        "sugestoes": "sugestões",
        "tecnicas": "técnicas",
        "familia": "família",
        "apresentacao": "apresentação",
        "condicoes": "condições",
        "programacao": "programação",
        "unico": "único",
        "instituicoes": "instituições",
        "varios": "vários",
        "dificuldades": "dificuldades",
        "rapido": "rápido",
        "atraves": "através",
        "operacoes": "operações",
        "calendario": "calendário",
        "Calendario": "Calendário",
        "utilizadores": "utilizadores",
        "administrador": "administrador",
        "gestor": "gestor",
        "motorista": "motorista",
        "botao": "botão",
        "rodape": "rodapé",
        "investigacao": "investigação",
        "esforco": "esforço",
        "evolucao": "evolução",
        "contribuiram": "contribuíram",
        "concretizacao": "concretização",
        "Agradeco": "Agradeço",
        "agradeco": "agradeço",
        "paciencia": "paciência",
        "conducao": "condução",
        "organizacao": "organização",
        "revisao": "revisão",
        "preparacao": "preparação",
        "manutencao": "manutenção",
        "Formulario": "Formulário",
        "formulario": "formulário",
        "Historico": "Histórico",
        "historico": "histórico",
        "Seguranca": "Segurança",
        "seguranca": "segurança",
        "Niveis": "Níveis",
        "niveis": "níveis",
        "Telemovel": "Telemóvel",
        "telemovel": "telemóvel",
        "Alcancados": "Alcançados",
        "alcancados": "alcançados",
        "Ecra": "Ecrã",
        "atualizar": "atualizar",
        "Apos": "Após",
        "apos": "após",
        "contem": "contém",
        "alem": "além",
        "tambem": "também",
        "esta": "está",
        "sao": "são",
        "nao": "não",
        "possivel": "possível",
        "disponivel": "disponível",
        "responsavel": "responsável",
        "proprios": "próprios",
        "necessario": "necessário",
        "Associacao": "Associação",
        "associacao": "associação",
        "Exportacao": "Exportação",
        "exportacao": "exportação",
        "Relacoes": "Relações",
        "relacoes": "relações",
        "Intervencoes": "Intervenções",
        "intervencoes": "intervenções",
        "Ocorrencias": "Ocorrências",
        "ocorrencias": "ocorrências",
        "operacional": "operacional",
        "estatisticos": "estatísticos",
        "notificacoes": "notificações",
        "prazos": "prazos",
        "moveis": "móveis",
    }

    def fix_text(text):
        for src, dst in replacements.items():
            if " " in src:
                text = text.replace(src, dst)
            else:
                text = re.sub(rf"\b{re.escape(src)}\b", dst, text)
        post_replacements = {
            "dàs": "das",
            "nàs": "nas",
            "pelàs": "pelas",
            "à gestão de frotas": "a gestão de frotas",
            "à gestão mais": "a gestão mais",
            "o Porque": "o Porquê",
            "O Porque": "O Porquê",
            "a desclassificação": "à desclassificação",
            "à escola": "à escola",
            "também a minha família": "também à minha família",
            "ligação a Base de Dados": "ligação à Base de Dados",
            "Ligação a Base de Dados": "Ligação à Base de Dados",
            "de gestao": "de gestão",
            "da gestao": "da gestão",
            "Sistema de gestao": "Sistema de gestão",
        }
        for src, dst in post_replacements.items():
            text = text.replace(src, dst)
        return text

    def fix_paragraph(paragraph):
        for run in paragraph.runs:
            run.text = fix_text(run.text)

    for paragraph in doc.paragraphs:
        fix_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fix_paragraph(paragraph)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Frotalink - Relatorio de Prova de Aptidao Profissional | Pagina ")
    set_run_font(r, size=9, color="555555")
    add_field(p, "PAGE")


def add_front_matter(doc):
    add_para(doc, "[Nome do Aluno]", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_para(doc, "Frotalink", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=30)
    add_para(doc, "Sistema de gestao de frota e operacao logistica", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True)
    add_para(doc, "Santo Tirso, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    page_break(doc)

    add_para(doc, "[Nome do Aluno]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Frotalink", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
    add_para(doc, "Relatorio de Prova de Aptidao Profissional apresentado a OFICINA - Escola Profissional do INA para a obtencao da Qualificacao Profissional de Nivel 4", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Curso Programador/a de Informatica", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Trabalho realizado sob a orientacao do professor [Nome do Orientador]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Junho, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    page_break(doc)

    add_heading(doc, "DECLARACAO", 1)
    add_table(
        doc,
        ["Elemento", "Informacao"],
        [
            ("Nome", "[Nome do Aluno]"),
            ("Endereco eletronico", "[email do aluno]"),
            ("Titulo do Relatorio da Prova de Aptidao Profissional", "Frotalink"),
            ("Orientador(es)", "[Nome do Orientador]"),
            ("Ano de Conclusao", "2026"),
            ("Designacao do Curso", "Programador/a de Informatica"),
        ],
        [5.0, 10.5],
    )
    add_para(doc, "E autorizada a reproducao integral deste relatorio da Prova de Aptidao Profissional apenas para efeitos de investigacao, mediante declaracao escrita do interessado, que a tal se compromete.")
    add_para(doc, "OFICINA - Escola Profissional do INA, ___/___/______")
    add_para(doc, "Assinatura: _______________________________________________")
    page_break(doc)

    add_heading(doc, "DECLARACAO DE COMPROMISSO ANTIPLAGIO", 1)
    add_para(doc, "Declaro por minha honra que o trabalho que apresento e original e que todas as citacoes, contribuicoes de textos e trabalhos alheios estao devidamente referenciados em notas de rodape e na bibliografia apresentada no final do trabalho.")
    add_para(doc, "Tenho consciencia de que a utilizacao de elementos alheios nao identificados constitui uma grave falta etica e disciplinar, podendo levar a desclassificacao da Prova de Aptidao Profissional.")
    add_para(doc, "_______________________________, ___ de ___________________ de 2026")
    add_para(doc, "______________________________________")
    add_para(doc, "Aluno/a")
    page_break(doc)

    add_heading(doc, "Agradecimentos", 1)
    add_para(doc, "A realizacao deste projeto representa o resultado de um percurso de aprendizagem, esforco e evolucao pessoal e profissional. Por esse motivo, quero agradecer a todos os que contribuiram, direta ou indiretamente, para a concretizacao da minha Prova de Aptidao Profissional.")
    add_para(doc, "Agradeco aos professores que acompanharam o desenvolvimento do projeto, pelo apoio, orientacao, disponibilidade e pelas sugestoes apresentadas ao longo das varias fases de trabalho. O acompanhamento recebido foi essencial para ultrapassar dificuldades tecnicas, melhorar a organizacao da aplicacao e consolidar o relatorio final.")
    add_para(doc, "Agradeco tambem a minha familia e aos meus colegas pelo incentivo, paciencia e apoio durante todo o processo. A sua ajuda foi importante nos momentos de maior trabalho, especialmente na revisao de ideias, preparacao da apresentacao e conclusao da documentacao.")
    add_para(doc, "Por fim, agradeco a escola pelas condicoes proporcionadas ao longo do curso, que permitiram adquirir conhecimentos nas areas de programacao, bases de dados, desenvolvimento web e documentacao tecnica.")
    page_break(doc)

    add_heading(doc, "Resumo", 1)
    add_para(doc, "O projeto Frotalink nasce da necessidade de tornar a gestao de frotas mais simples, organizada e eficiente. Muitas empresas e instituicoes que utilizam varios veiculos no seu dia a dia enfrentam dificuldades em acompanhar o estado da frota, controlar manutencoes, gerir motoristas, consultar informacoes de clientes, organizar cargas e manter todos os dados atualizados num unico local.")
    add_para(doc, "Ao analisar este problema, concluiu-se que seria util criar uma plataforma digital capaz de centralizar a informacao operacional relacionada com a frota. Assim, o Frotalink pretende reduzir falhas de comunicacao, evitar esquecimentos associados a revisoes, avarias ou abastecimentos e permitir um controlo mais rapido dos veiculos, motoristas e cargas.")
    add_para(doc, "O Frotalink consiste no desenvolvimento de uma aplicacao web de gestao de frota e operacao logistica. A plataforma permite consultar, inserir, editar e organizar dados sobre viaturas, motoristas, clientes, cargas, manutencoes, avarias, abastecimentos, utilizadores, relatorios e calendario. A aplicacao foi desenvolvida com PHP, MySQL/MariaDB, HTML, CSS e JavaScript, funcionando em ambiente local atraves do XAMPP.")
    page_break(doc)

    add_heading(doc, "Palavras-chave", 1)
    add_para(doc, "frotalink, gestao de frotas, veiculos, viaturas, motoristas, condutores, manutencao, avarias, abastecimentos, combustivel, clientes, cargas, logistica, transportes, calendario, relatorios, pesquisa, utilizadores, administrador, gestor, motorista, programacao, web, webserver, frontend, backend, base de dados, PHP, MySQL, MariaDB, SQL, HTML, CSS, JavaScript, XAMPP.")
    page_break(doc)

    add_heading(doc, "Indice de Figuras", 1)
    figures = [
        "Figura 1 - Logotipo do Frotalink",
        "Figura 2 - Ferramentas utilizadas",
        "Figura 3 - Fluxo de informacao do sistema",
        "Figura 4 - Diagrama simplificado da base de dados",
        "Figura 5 - Arquitetura geral da plataforma",
        "Figura 6 - Estrutura de pastas do projeto",
        "Figura 7 - Ficheiro config.php",
        "Figura 8 - Pagina de login",
        "Figura 9 - Dashboard principal",
        "Figura 10 - Menu lateral",
        "Figura 11 - Gestao de viaturas",
        "Figura 12 - Formulario de criacao de viatura",
        "Figura 13 - Historico da viatura",
        "Figura 14 - Gestao de motoristas",
        "Figura 15 - Gestao de clientes",
        "Figura 16 - Gestao de cargas",
        "Figura 17 - Formulario de criacao de carga",
        "Figura 18 - Gestao de manutencoes",
        "Figura 19 - Gestao de avarias e problemas",
        "Figura 20 - Gestao de abastecimentos",
        "Figura 21 - Calendario",
        "Figura 22 - Pesquisa global",
        "Figura 23 - Relatorios",
        "Figura 24 - Gestao de utilizadores",
        "Figura 25 - Codigo de ligacao a base de dados",
        "Figura 26 - Codigo de controlo de sessoes",
        "Figura 27 - Codigo de protecao CSRF",
        "Figura 28 - Tabelas da base de dados",
        "Figura 29 - Teste de login",
        "Figura 30 - Teste de criacao de registos",
        "Figura 31 - Teste de edicao de registos",
        "Figura 32 - Teste de eliminacao de registos",
        "Figura 33 - Teste de responsividade em computador",
        "Figura 34 - Teste de responsividade em telemovel",
    ]
    for f in figures:
        add_para(doc, f + " ........................................ pag. ___")
    page_break(doc)

    add_heading(doc, "Indice", 1)
    toc_entries = [
        ("Conteudo", ""),
        ("Resumo", "VI"),
        ("Palavras-chave", "VII"),
        ("Indice De Figuras", "VIII"),
        ("1. Introducao", "1"),
        ("1.1 Proposito", "2"),
        ("1.2 Ambito", "2"),
        ("1.3 Definicoes, Acronimos e Abreviaturas", "3"),
        ("1.4 Visao Geral Do Presente Documento", "4"),
        ("2. Anteprojeto Da Prova De Aptidao Profissional", "5"),
        ("2.1 Identificacao do Projeto", "5"),
        ("2.2 Problema Identificado", "6"),
        ("2.3 Objetivos do Projeto", "7"),
        ("2.4 Publico-Alvo", "8"),
        ("2.5 Recursos Necessarios", "8"),
        ("2.6 Planeamento Inicial", "9"),
        ("3. Desenvolvimento do Projeto", "10"),
        ("3.1 Como Surgiu a Ideia", "10"),
        ("3.2 Logotipo", "11"),
        ("3.3 Ferramentas Utilizadas", "12"),
        ("3.4 Fluxo de Informacao", "14"),
        ("3.5 Arquitetura do Sistema", "15"),
        ("3.6 Estrutura de Pastas do Projeto", "16"),
        ("3.7 Documentacao Tecnica", "17"),
        ("4. Modulos da Aplicacao", "24"),
        ("5. Interfaces da Plataforma", "38"),
        ("6. Base de Dados", "56"),
        ("7. Seguranca", "64"),
        ("8. Testes Realizados", "70"),
        ("9. Apreciacao Critica", "78"),
        ("10. Bibliografia", "82"),
        ("11. Anexos", "84"),
    ]
    for label, page in toc_entries:
        dots = "." * max(8, 70 - len(label))
        add_para(doc, f"{label} {dots} {page}" if page else label)
    page_break(doc)


def add_main_content(doc):
    add_heading(doc, "1. Introducao", 1)
    add_para(doc, "O presente projeto, denominado Frotalink, foi desenvolvido no ambito da Prova de Aptidao Profissional do curso de Programador/a de Informatica. A aplicacao tem como objetivo apoiar a gestao de frotas e a organizacao de processos logisticos atraves de uma plataforma web centralizada.")
    add_para(doc, "A gestao de uma frota implica acompanhar veiculos, motoristas, clientes, cargas, manutencoes, avarias e abastecimentos. Quando estas informacoes se encontram espalhadas por folhas de calculo, documentos em papel ou mensagens, torna-se dificil obter uma visao clara do estado da operacao. Esta dificuldade pode originar atrasos, esquecimentos, duplicacao de informacao e perda de controlo sobre recursos importantes.")
    add_para(doc, "O Frotalink surge como resposta a esta necessidade. A plataforma permite organizar dados essenciais da frota num unico sistema, com autenticacao de utilizadores, niveis de acesso e modulos dedicados as principais areas de trabalho. Desta forma, o projeto procura facilitar a consulta de informacao, melhorar a organizacao interna e tornar os processos de gestao mais eficientes.")

    add_heading(doc, "1.1 Proposito", 2)
    add_para(doc, "Este documento tem como proposito apresentar de forma detalhada o desenvolvimento do projeto Frotalink, descrevendo o problema identificado, os objetivos definidos, as ferramentas utilizadas, a estrutura tecnica, a base de dados, as interfaces criadas, os testes realizados e a avaliacao final do trabalho desenvolvido.")

    add_heading(doc, "1.2 Ambito", 2)
    add_para(doc, "O ambito do projeto centra-se no desenvolvimento de uma aplicacao web para gestao de frota e operacao logistica. A aplicacao foi pensada para ser utilizada por administradores, gestores e motoristas, permitindo que cada perfil tenha acesso as funcionalidades adequadas a sua funcao.")
    add_bullets(doc, [
        "Administradores podem gerir utilizadores e aceder as principais areas da plataforma.",
        "Gestores podem acompanhar a operacao, gerir viaturas, motoristas, clientes, cargas, manutencoes, avarias e abastecimentos.",
        "Motoristas podem consultar informacao operacional associada ao seu trabalho, especialmente cargas e ocorrencias.",
    ])

    add_heading(doc, "1.3 Definicoes, Acronimos e Abreviaturas", 2)
    add_table(doc, ["Termo", "Definicao"], [
        ("Frota", "Conjunto de veiculos utilizados por uma organizacao."),
        ("Viatura", "Veiculo registado na plataforma para consulta e gestao."),
        ("Motorista", "Utilizador ou colaborador responsavel pela conducao e operacao de uma viatura."),
        ("CRUD", "Conjunto de operacoes Create, Read, Update e Delete: criar, consultar, atualizar e eliminar."),
        ("Backend", "Parte do sistema responsavel pela logica, ligacao a base de dados e processamento de pedidos."),
        ("Frontend", "Parte visual da aplicacao com a qual o utilizador interage."),
        ("PHP", "Linguagem utilizada no desenvolvimento da logica da aplicacao web."),
        ("MySQL/MariaDB", "Sistema de gestao de base de dados utilizado para guardar informacao."),
        ("XAMPP", "Ambiente local que inclui Apache, PHP e MySQL/MariaDB."),
        ("CSRF", "Tipo de ataque em formularios web; no projeto e mitigado atraves de tokens de seguranca."),
    ], [4.0, 11.5])

    add_heading(doc, "1.4 Visao Geral do Documento", 2)
    add_para(doc, "Depois da introducao, o documento apresenta o anteprojeto da PAP, onde sao descritos o problema, os objetivos, o publico-alvo e o planeamento. Em seguida, e explicado o desenvolvimento do projeto, incluindo a ideia, as ferramentas utilizadas, a arquitetura, a base de dados e a documentacao tecnica.")
    add_para(doc, "Posteriormente, sao detalhados os modulos da aplicacao, as interfaces, a seguranca implementada e os testes realizados. No final, apresenta-se uma apreciacao critica sobre o trabalho desenvolvido, as dificuldades encontradas, os objetivos alcancados e possiveis melhorias futuras.")

    add_heading(doc, "2. Anteprojeto da Prova de Aptidao Profissional", 1)
    add_heading(doc, "2.1 Identificacao do Projeto", 2)
    add_table(doc, ["Elemento", "Descricao"], [
        ("Nome do projeto", "Frotalink"),
        ("Tema", "Sistema de gestao de frota e operacao logistica"),
        ("Area", "Desenvolvimento web, bases de dados e gestao operacional"),
        ("Tecnologias", "PHP, MySQL/MariaDB, HTML, CSS, JavaScript e XAMPP"),
        ("Objetivo principal", "Centralizar e simplificar a gestao de viaturas, motoristas, clientes, cargas e manutencoes."),
    ], [4.5, 11.0])

    add_heading(doc, "2.2 Problema Identificado", 2)
    add_para(doc, "O problema identificado esta relacionado com a dificuldade de organizar informacao operacional numa frota. Em muitas situacoes, os dados sobre viaturas, motoristas, cargas, clientes, manutencoes e avarias ficam dispersos por diferentes meios, o que dificulta a consulta e o acompanhamento do estado real da operacao.")
    add_para(doc, "Esta falta de centralizacao pode provocar atrasos, esquecimentos de manutencoes, dificuldades em associar cargas a motoristas e veiculos, perda de historico e pouca clareza sobre custos de combustivel ou ocorrencias. Assim, tornou-se pertinente desenvolver uma solucao simples, acessivel e adaptada a realidade de uma pequena ou media operacao de transportes.")

    add_heading(doc, "2.3 Objetivos do Projeto", 2)
    add_para(doc, "O objetivo geral do Frotalink e desenvolver uma aplicacao web funcional para apoiar a gestao de frota e operacoes de transporte.")
    add_bullets(doc, [
        "Implementar autenticacao de utilizadores.",
        "Criar niveis de acesso para administrador, gestor e motorista.",
        "Permitir o registo, consulta, edicao e eliminacao de viaturas.",
        "Permitir a gestao de motoristas e a sua associacao a utilizadores.",
        "Permitir o registo e acompanhamento de clientes e cargas.",
        "Registar manutencoes, avarias, abastecimentos e historico operacional.",
        "Disponibilizar dashboard, calendario, pesquisa global e relatorios.",
        "Construir uma interface simples, organizada e adequada ao uso diario.",
    ])

    add_heading(doc, "2.4 Publico-Alvo", 2)
    add_para(doc, "O publico-alvo do Frotalink inclui pequenas empresas, equipas ou instituicoes que necessitam de acompanhar uma frota de veiculos e organizar processos logisticos. A plataforma tambem pode ser utilizada em contexto escolar como demonstracao de uma aplicacao web com base de dados, autenticacao e modulos CRUD.")

    add_heading(doc, "2.5 Recursos Necessarios", 2)
    add_table(doc, ["Recurso", "Utilizacao"], [
        ("Computador", "Desenvolvimento, testes e execucao da aplicacao."),
        ("XAMPP", "Servidor local Apache e base de dados MySQL/MariaDB."),
        ("Editor de codigo", "Criacao e alteracao dos ficheiros PHP, CSS e JavaScript."),
        ("Navegador web", "Teste das paginas e validacao da interface."),
        ("phpMyAdmin", "Consulta e validacao da base de dados."),
        ("Word", "Edicao e finalizacao do relatorio da PAP."),
    ], [4.0, 11.5])

    add_heading(doc, "2.6 Planeamento Inicial", 2)
    add_numbered(doc, [
        "Identificacao do problema e definicao do tema.",
        "Levantamento das entidades principais da base de dados.",
        "Criacao da estrutura inicial do projeto em PHP.",
        "Implementacao da autenticacao e das sessoes.",
        "Desenvolvimento dos modulos CRUD.",
        "Criacao do dashboard, pesquisa, calendario e relatorios.",
        "Testes aos fluxos principais.",
        "Preparacao da documentacao e apresentacao final.",
    ])

    add_heading(doc, "3. Desenvolvimento do Projeto", 1)
    add_heading(doc, "3.1 Como Surgiu a Ideia", 2)
    add_para(doc, "A ideia do Frotalink surgiu da observacao de uma necessidade comum em contextos onde existem varios veiculos e pessoas envolvidas na sua utilizacao. Para gerir uma frota, nao basta saber quantos veiculos existem; e necessario acompanhar o seu estado, saber quem os utiliza, que cargas estao associadas, que manutencoes existem e que problemas foram reportados.")
    add_para(doc, "Ao refletir sobre esta realidade, concluiu-se que uma aplicacao web poderia facilitar muito este processo, permitindo reunir toda a informacao num unico local. A escolha por uma aplicacao web deve-se ao facto de poder ser utilizada num navegador, sem instalacao complexa, e por permitir uma estrutura simples de demonstrar e testar em ambiente local.")

    add_heading(doc, "3.2 Logotipo", 2)
    add_para(doc, "O logotipo do Frotalink representa a identidade visual do projeto. O nome junta a ideia de frota com a ideia de ligacao, sugerindo uma plataforma que conecta veiculos, motoristas, clientes e operacoes.")
    add_figure_placeholder(doc, "Figura 1 - Logotipo do Frotalink", "Inserir o ficheiro assets/logo.svg ou uma imagem exportada do logotipo.")

    add_heading(doc, "3.3 Ferramentas Utilizadas", 2)
    add_table(doc, ["Ferramenta", "Funcao no projeto"], [
        ("XAMPP", "Execucao local do servidor Apache e da base de dados."),
        ("PHP", "Desenvolvimento do backend e das paginas dinamicas."),
        ("MySQL/MariaDB", "Armazenamento de dados da aplicacao."),
        ("HTML", "Estrutura das paginas web."),
        ("CSS", "Estilizacao visual da interface."),
        ("JavaScript", "Interacoes simples no frontend, como pesquisa no topo."),
        ("phpMyAdmin", "Gestao e verificacao da base de dados."),
        ("Navegador web", "Teste da interface e dos fluxos de utilizacao."),
    ], [4.0, 11.5])
    add_figure_placeholder(doc, "Figura 2 - Ferramentas utilizadas")

    add_heading(doc, "3.4 Linguagens Utilizadas", 2)
    add_para(doc, "As linguagens utilizadas no Frotalink foram escolhidas tendo em conta o objetivo de construir uma aplicacao web funcional, simples de executar em ambiente local e adequada aos conhecimentos desenvolvidos ao longo do curso.")
    add_bullets(doc, [
        "PHP, para processar formularios, validar sessoes, comunicar com a base de dados e gerar paginas dinamicas.",
        "SQL, para criacao, consulta e atualizacao dos dados na base de dados.",
        "HTML, para estruturar os formularios, listas, tabelas e conteudos das paginas.",
        "CSS, para definir o aspeto visual da plataforma.",
        "JavaScript, para pequenas funcionalidades de interacao no navegador.",
    ])

    add_heading(doc, "3.5 O Porque da Escolha das Tecnologias", 2)
    add_para(doc, "A escolha de PHP e MySQL/MariaDB deveu-se a sua forte ligacao ao desenvolvimento web tradicional, a facilidade de configuracao com XAMPP e a possibilidade de criar rapidamente paginas dinamicas com acesso a base de dados. Estas tecnologias sao adequadas para projetos escolares e para aplicacoes internas simples.")
    add_para(doc, "HTML, CSS e JavaScript foram utilizados por serem tecnologias base da web. O HTML define a estrutura, o CSS organiza a apresentacao visual e o JavaScript permite melhorar a experiencia do utilizador em pequenas interacoes.")

    add_heading(doc, "3.6 Desafios", 2)
    add_para(doc, "Durante o desenvolvimento surgiram varios desafios, especialmente relacionados com a organizacao da base de dados, validacao de formularios, controlo de permissoes e ligacao entre entidades. Tambem foi necessario garantir que as paginas mantinham uma estrutura visual coerente e que os diferentes modulos funcionavam de forma semelhante.")
    add_bullets(doc, [
        "Definir corretamente as tabelas e as relacoes entre dados.",
        "Evitar acessos indevidos atraves de sessoes e niveis de permissao.",
        "Criar formularios completos mas simples de utilizar.",
        "Manter coerencia entre paginas de listagem, criacao, edicao e eliminacao.",
        "Preparar dados de demonstracao para apresentar o projeto.",
    ])

    add_heading(doc, "3.7 Fluxo de Informacao", 2)
    add_para(doc, "O fluxo de informacao do Frotalink comeca no utilizador, que acede a uma pagina da plataforma atraves do navegador. O pedido e processado por ficheiros PHP, que verificam a sessao, validam permissoes e comunicam com a base de dados MySQL/MariaDB. A resposta e apresentada novamente ao utilizador atraves de paginas HTML estilizadas com CSS.")
    add_figure_placeholder(doc, "Figura 3 - Fluxo de informacao do sistema")

    add_heading(doc, "3.8 Base de Dados Utilizada", 2)
    add_para(doc, "A base de dados utilizada chama-se gestao_frotas. A estrutura e criada automaticamente no ficheiro config.php atraves de instrucoes CREATE TABLE IF NOT EXISTS. Esta abordagem facilita a preparacao do projeto, pois garante que as tabelas essenciais existem quando a aplicacao e executada.")
    add_figure_placeholder(doc, "Figura 4 - Diagrama simplificado da base de dados")

    add_heading(doc, "3.9 Arquitetura do Sistema", 2)
    add_para(doc, "A arquitetura do Frotalink segue uma organizacao simples baseada em paginas PHP. Cada modulo possui ficheiros proprios para listar, criar, editar e apagar registos. O ficheiro config.php centraliza a ligacao a base de dados, a criacao inicial das tabelas, as funcoes de sessao, permissoes e seguranca.")
    add_figure_placeholder(doc, "Figura 5 - Arquitetura geral da plataforma")

    add_heading(doc, "3.10 Documentacao Tecnica", 2)
    add_para(doc, "A documentacao tecnica do projeto inclui a descricao da estrutura de ficheiros, a explicacao das tabelas da base de dados, a identificacao dos modulos implementados e a demonstracao das principais funcionalidades. Esta documentacao permite compreender como a aplicacao esta organizada e como pode ser mantida ou evoluida no futuro.")

    add_heading(doc, "4. Funcionamento Tecnico do Frotalink", 1)
    add_heading(doc, "4.1 Estrutura de Pastas do Projeto", 2)
    add_para(doc, "O projeto esta organizado numa pasta principal chamada frotas. Dentro desta pasta encontram-se os ficheiros PHP da aplicacao, a folha de estilos, os recursos visuais e a documentacao.")
    add_table(doc, ["Ficheiro/Pasta", "Descricao"], [
        ("assets/", "Recursos da interface, como logotipo e JavaScript auxiliar."),
        ("docs/", "Documentos de apoio e scripts relacionados com a documentacao."),
        ("entregas_pap/", "Ficheiros preparados para entrega da PAP."),
        ("config.php", "Configuracao, ligacao a base de dados, sessoes, permissoes e CSRF."),
        ("layout.css", "Estilos visuais da aplicacao."),
        ("index.php", "Dashboard principal apos login."),
        ("login.php / logout.php", "Entrada e saida do sistema."),
        ("*_listar.php", "Paginas de listagem dos varios modulos."),
        ("*_criar.php", "Paginas de criacao de registos."),
        ("*_editar.php", "Paginas de edicao de registos."),
        ("*_apagar.php", "Paginas de eliminacao de registos."),
    ], [5.0, 10.5])
    add_figure_placeholder(doc, "Figura 6 - Estrutura de pastas do projeto")

    add_heading(doc, "4.2 Ficheiro config.php", 2)
    add_para(doc, "O ficheiro config.php e um dos ficheiros mais importantes do projeto. Nele encontram-se as constantes de ligacao a base de dados, a funcao getDbConnection, a criacao das tabelas, as funcoes de sessao, os niveis de acesso e os mecanismos de protecao CSRF.")
    add_code(doc, "define('DB_HOST', 'localhost');\ndefine('DB_USER', 'root');\ndefine('DB_NAME', 'gestao_frotas');")
    add_figure_placeholder(doc, "Figura 7 - Ficheiro config.php")

    add_heading(doc, "4.3 Ligacao a Base de Dados", 2)
    add_para(doc, "A ligacao a base de dados e feita atraves da extensao mysqli. A funcao getDbConnection cria uma nova ligacao, verifica erros de conexao e define o conjunto de caracteres utf8mb4, garantindo suporte para texto em portugues e caracteres especiais.")
    add_code(doc, "function getDbConnection() {\n    $mysqli = new mysqli(DB_HOST, DB_USER, DB_PASS, DB_NAME);\n    $mysqli->set_charset(DB_CHARSET);\n    return $mysqli;\n}")
    add_figure_placeholder(doc, "Figura 25 - Codigo de ligacao a base de dados")

    add_heading(doc, "4.4 Sistema de Login e Sessoes", 2)
    add_para(doc, "O sistema de login permite autenticar utilizadores e guardar na sessao informacao essencial como identificador, nome e nivel de acesso. As paginas protegidas chamam funcoes que verificam se o utilizador esta autenticado antes de permitir o acesso.")
    add_code(doc, "function isLoggedIn() {\n    return !empty($_SESSION['user']) && !empty($_SESSION['user_id']) && !empty($_SESSION['nivel_acesso']);\n}")
    add_figure_placeholder(doc, "Figura 26 - Codigo de controlo de sessoes")

    add_heading(doc, "4.5 Controlo de Permissoes", 2)
    add_para(doc, "O controlo de permissoes permite separar funcionalidades entre administrador, gestor e motorista. Esta separacao e importante para impedir que utilizadores sem autorizacao acedam a areas sensiveis, como a gestao de utilizadores.")
    add_table(doc, ["Perfil", "Permissoes principais"], [
        ("Administrador", "Acesso geral, incluindo gestao de utilizadores."),
        ("Gestor", "Acesso a operacao: viaturas, motoristas, clientes, cargas, manutencoes, avarias e relatorios."),
        ("Motorista", "Acesso limitado as funcionalidades relacionadas com a sua operacao."),
    ], [4.0, 11.5])

    add_heading(doc, "4.6 Protecao com CSRF", 2)
    add_para(doc, "A protecao CSRF foi implementada com tokens guardados em sessao. Cada formulario pode incluir um campo escondido com o token, e o servidor valida se o token recebido corresponde ao token guardado. Esta medida ajuda a proteger operacoes sensiveis contra pedidos nao autorizados.")
    add_code(doc, "function verify_csrf_token($token) {\n    return !empty($token) && !empty($_SESSION['csrf_token']) && hash_equals($_SESSION['csrf_token'], $token);\n}")
    add_figure_placeholder(doc, "Figura 27 - Codigo de protecao CSRF")

    add_heading(doc, "4.7 Gestao de Utilizadores", 2)
    add_para(doc, "A gestao de utilizadores permite criar, listar, editar e remover contas da plataforma. Cada utilizador possui nome, email, palavra-passe encriptada, nivel de acesso e estado. Este modulo e essencial para controlar quem pode entrar no sistema e que funcionalidades pode utilizar.")

    add_heading(doc, "4.8 Gestao de Viaturas", 2)
    add_para(doc, "O modulo de viaturas permite registar veiculos com matricula, modelo, estado, quilometragem total, consumo medio e localizacao. A gestao de viaturas e uma das bases do Frotalink, pois muitos outros modulos dependem da existencia de veiculos registados.")

    add_heading(doc, "4.9 Gestao de Motoristas", 2)
    add_para(doc, "O modulo de motoristas permite guardar dados pessoais e profissionais, como nome, contacto, carta de conducao, validade da carta, categoria, numero mecanografico, tipo de contrato, estado, disponibilidade e viatura atual. Tambem permite associar um motorista a uma conta de utilizador.")

    add_heading(doc, "4.10 Gestao de Clientes", 2)
    add_para(doc, "O modulo de clientes permite registar pessoas ou entidades juridicas que estejam associadas a servicos de transporte. Inclui dados como nome, NIF/NIPC, morada, contacto, email, limite de credito, prazo de pagamento, estado e coordenadas geograficas quando disponiveis.")

    add_heading(doc, "4.11 Gestao de Cargas", 2)
    add_para(doc, "O modulo de cargas permite acompanhar transportes desde a fase pendente ate a entrega. Cada carga pode ter codigo de rastreio, estado, descricao, tipo, peso, volume, paletes, locais e datas de recolha e entrega, cliente, valor de transporte, pagamento, viatura e motorista associados.")

    add_heading(doc, "4.12 Gestao de Manutencoes", 2)
    add_para(doc, "O modulo de manutencoes permite registar intervencoes associadas aos veiculos. Inclui tipo, estado, criticidade, descricao e data agendada. Este modulo ajuda a evitar esquecimentos e a manter uma melhor organizacao das intervencoes tecnicas.")

    add_heading(doc, "4.13 Gestao de Avarias", 2)
    add_para(doc, "A area de avarias e problemas permite registar ocorrencias relacionadas com viaturas ou cargas. Cada registo possui titulo, descricao, prioridade, estado, utilizador que reportou e possivel resposta do gestor. Esta funcionalidade facilita a comunicacao de problemas e o seu acompanhamento.")

    add_heading(doc, "4.14 Gestao de Abastecimentos", 2)
    add_para(doc, "O modulo de abastecimentos permite registar consumos de combustivel, incluindo veiculo, motorista, data, litros, custo total, odometro, posto e observacoes. Estes dados permitem acompanhar custos e consumos associados a frota.")

    add_heading(doc, "4.15 Calendario", 2)
    add_para(doc, "O calendario permite visualizar eventos e datas importantes relacionadas com a operacao, como manutencoes, viagens ou tarefas planeadas. Esta area ajuda a ter uma visao temporal da atividade da frota.")

    add_heading(doc, "4.16 Relatorios", 2)
    add_para(doc, "A pagina de relatorios permite consultar informacao resumida e exportar dados relevantes. Esta funcionalidade e importante para apoiar a analise da operacao e apresentar resultados de forma organizada.")

    add_heading(doc, "4.17 Pesquisa Global", 2)
    add_para(doc, "A pesquisa global permite procurar rapidamente informacao dentro da plataforma. Esta funcionalidade melhora a usabilidade, pois reduz o tempo necessario para encontrar registos especificos.")

    add_heading(doc, "5. Interfaces da Plataforma", 1)
    interfaces = [
        ("5.1 Pagina de Login", "A pagina de login permite aos utilizadores autenticarem-se com as suas credenciais.", "Figura 8 - Pagina de login"),
        ("5.2 Pagina Inicial / Dashboard", "O dashboard apresenta uma visao geral da operacao, com indicadores e atalhos para os modulos principais.", "Figura 9 - Dashboard principal"),
        ("5.3 Menu Lateral", "O menu lateral organiza a navegacao da plataforma e facilita o acesso as diferentes areas.", "Figura 10 - Menu lateral"),
        ("5.4 Interface de Viaturas", "A interface de viaturas permite consultar a lista de veiculos e aceder as opcoes de criacao, edicao, eliminacao e historico.", "Figura 11 - Gestao de viaturas"),
        ("5.5 Formulario de Viatura", "O formulario de viatura recolhe dados como matricula, modelo, estado, quilometragem e consumo medio.", "Figura 12 - Formulario de criacao de viatura"),
        ("5.6 Historico da Viatura", "O historico da viatura permite acompanhar informacao associada ao veiculo ao longo do tempo.", "Figura 13 - Historico da viatura"),
        ("5.7 Interface de Motoristas", "A interface de motoristas organiza os dados dos condutores e a sua disponibilidade.", "Figura 14 - Gestao de motoristas"),
        ("5.8 Interface de Clientes", "A interface de clientes permite gerir contactos, dados fiscais, moradas e estados dos clientes.", "Figura 15 - Gestao de clientes"),
        ("5.9 Interface de Cargas", "A interface de cargas permite acompanhar o estado dos transportes e as associacoes a viaturas, motoristas e clientes.", "Figura 16 - Gestao de cargas"),
        ("5.10 Formulario de Carga", "O formulario de carga permite registar os dados operacionais de recolha, entrega, pagamento e transporte.", "Figura 17 - Formulario de criacao de carga"),
        ("5.11 Interface de Manutencoes", "A interface de manutencoes permite planear e acompanhar intervencoes nos veiculos.", "Figura 18 - Gestao de manutencoes"),
        ("5.12 Interface de Avarias", "A interface de avarias centraliza problemas reportados e respostas de gestao.", "Figura 19 - Gestao de avarias e problemas"),
        ("5.13 Interface de Combustivel", "A interface de combustivel apresenta registos de abastecimentos e custos associados.", "Figura 20 - Gestao de abastecimentos"),
        ("5.14 Interface de Calendario", "O calendario ajuda a visualizar datas relevantes de forma organizada.", "Figura 21 - Calendario"),
        ("5.15 Interface de Pesquisa", "A pesquisa global permite encontrar rapidamente registos em diferentes areas.", "Figura 22 - Pesquisa global"),
        ("5.16 Interface de Relatorios", "A area de relatorios apoia a consulta e exportacao de dados operacionais.", "Figura 23 - Relatorios"),
        ("5.17 Interface de Utilizadores", "A gestao de utilizadores permite controlar contas, niveis de acesso e estado.", "Figura 24 - Gestao de utilizadores"),
    ]
    for heading, text, fig in interfaces:
        add_heading(doc, heading, 2)
        add_para(doc, text)
        add_figure_placeholder(doc, fig)

    add_heading(doc, "6. Base de Dados", 1)
    add_para(doc, "A base de dados do Frotalink foi criada para suportar os principais processos da aplicacao. As tabelas representam entidades reais do dominio da gestao de frotas, como veiculos, motoristas, clientes, cargas, manutencoes, avarias e abastecimentos.")
    add_figure_placeholder(doc, "Figura 28 - Tabelas da base de dados")
    add_table(doc, ["Tabela", "Finalidade"], [
        ("users", "Tabela inicial de utilizadores simples, incluindo utilizador administrador por defeito."),
        ("utilizadores", "Contas principais da plataforma, com email, password hash, nivel de acesso e estado."),
        ("veiculos", "Registo das viaturas, matriculas, modelos, estado, quilometragem e localizacao."),
        ("viagens", "Registo de viagens associadas a um veiculo e distancia percorrida."),
        ("manutencoes", "Intervencoes agendadas ou concluidas sobre veiculos."),
        ("historico_manutencoes_inspecoes", "Historico tecnico detalhado de inspecoes e manutencoes."),
        ("motoristas", "Dados pessoais, profissionais e operacionais dos motoristas."),
        ("clientes", "Dados comerciais e de contacto dos clientes."),
        ("cargas", "Operacoes de transporte, recolha, entrega, pagamento, motorista e viatura."),
        ("avarias_problemas", "Ocorrencias reportadas e acompanhamento de problemas."),
        ("abastecimentos", "Registos de combustivel, custos, litros e odometro."),
    ], [5.0, 10.5])

    add_heading(doc, "6.1 Relacoes entre Tabelas", 2)
    add_para(doc, "As relacoes entre tabelas sao estabelecidas atraves de chaves estrangeiras. Por exemplo, uma carga pode estar associada a uma viatura e a um motorista; uma manutencao esta associada a um veiculo; um abastecimento pertence a um veiculo e pode estar associado a um motorista; uma avaria pode estar relacionada com uma viatura, uma carga e utilizadores que reportam ou resolvem o problema.")

    add_heading(doc, "6.2 Principais Campos por Entidade", 2)
    add_table(doc, ["Entidade", "Campos principais"], [
        ("Viaturas", "matricula, modelo, status, km_total, consumo_medio, lat, lon"),
        ("Motoristas", "nome_completo, telefone, email, carta de conducao, estado, disponibilidade"),
        ("Clientes", "tipo_cliente, nome, nif_nipc, morada, telefone, email, estado_cliente"),
        ("Cargas", "codigo_rastreio, estado_carga, peso, locais, datas, cliente, viatura, motorista"),
        ("Manutencoes", "veiculo_id, tipo, status, criticidade, descricao, data_agendada"),
        ("Avarias", "titulo, descricao, prioridade, status, viatura, carga, reportado_por"),
        ("Abastecimentos", "veiculo, motorista, data, litros, custo_total, odometro"),
    ], [4.0, 11.5])

    add_heading(doc, "7. Seguranca", 1)
    add_heading(doc, "7.1 Autenticacao de Utilizadores", 2)
    add_para(doc, "A autenticacao garante que apenas utilizadores registados conseguem aceder a plataforma. O sistema verifica as credenciais e cria uma sessao quando o login e valido.")

    add_heading(doc, "7.2 Palavras-passe Encriptadas", 2)
    add_para(doc, "As palavras-passe sao guardadas com password_hash, evitando que fiquem armazenadas em texto simples na base de dados. Esta pratica e essencial para proteger os dados dos utilizadores.")

    add_heading(doc, "7.3 Niveis de Acesso", 2)
    add_para(doc, "Os niveis de acesso permitem limitar funcionalidades de acordo com o perfil. Esta separacao reduz riscos e torna a plataforma mais organizada.")

    add_heading(doc, "7.4 Protecao Contra Acessos Indevidos", 2)
    add_para(doc, "Funcoes como requireLogin, requireRole, requireOperationsAccess e requireUserManagementAccess garantem que as paginas protegidas apenas sao carregadas por utilizadores autorizados.")

    add_heading(doc, "7.5 Validacao de Formularios", 2)
    add_para(doc, "A validacao de formularios ajuda a evitar dados incompletos ou incorretos. Em operacoes sensiveis, a protecao CSRF acrescenta uma camada adicional de seguranca.")

    add_heading(doc, "8. Testes Realizados", 1)
    add_para(doc, "Os testes realizados tiveram como objetivo verificar se os principais fluxos da plataforma funcionavam corretamente. Foram testados login, criacao, edicao e eliminacao de registos, permissao de acessos e consulta de informacao.")
    tests = [
        ("8.1 Teste ao Login", "Foi testada a autenticacao com credenciais validas e invalidas, verificando se o sistema permite ou bloqueia corretamente o acesso.", "Figura 29 - Teste de login"),
        ("8.2 Teste a Criacao de Registos", "Foram criados registos em modulos como viaturas, clientes, motoristas, cargas e manutencoes.", "Figura 30 - Teste de criacao de registos"),
        ("8.3 Teste a Edicao de Dados", "Foram alterados dados existentes para confirmar se as atualizacoes eram guardadas na base de dados.", "Figura 31 - Teste de edicao de registos"),
        ("8.4 Teste a Eliminacao de Registos", "Foram testadas operacoes de remocao, respeitando dependencias e relacoes entre tabelas.", "Figura 32 - Teste de eliminacao de registos"),
        ("8.5 Teste de Responsividade em Computador", "Foi verificado se a interface se apresenta corretamente num ecra de computador.", "Figura 33 - Teste de responsividade em computador"),
        ("8.6 Teste de Responsividade em Telemovel", "Foi verificado se as paginas continuam utilizaveis em ecra reduzido.", "Figura 34 - Teste de responsividade em telemovel"),
    ]
    for h, t, f in tests:
        add_heading(doc, h, 2)
        add_para(doc, t)
        add_figure_placeholder(doc, f)

    add_heading(doc, "8.7 Resultados dos Testes", 2)
    add_table(doc, ["Area testada", "Resultado esperado", "Estado"], [
        ("Login", "Utilizador valido entra na plataforma; utilizador invalido e rejeitado.", "Aprovado"),
        ("Viaturas", "Permitir criar, listar, editar, apagar e consultar historico.", "Aprovado"),
        ("Motoristas", "Permitir gerir dados e disponibilidade.", "Aprovado"),
        ("Clientes", "Permitir gerir dados comerciais e contactos.", "Aprovado"),
        ("Cargas", "Permitir acompanhar estados e associacoes operacionais.", "Aprovado"),
        ("Manutencoes", "Permitir registar e acompanhar intervencoes.", "Aprovado"),
        ("Avarias", "Permitir reportar e acompanhar problemas.", "Aprovado"),
        ("Abastecimentos", "Permitir registar custos e consumos.", "Aprovado"),
        ("Permissoes", "Bloquear acessos a areas nao autorizadas.", "Aprovado"),
    ], [4.0, 8.0, 3.5])

    add_heading(doc, "9. Apreciacao Critica", 1)
    add_heading(doc, "9.1 Objetivos Alcancados", 2)
    add_para(doc, "O projeto conseguiu atingir os principais objetivos definidos. Foi desenvolvida uma aplicacao web funcional, com autenticacao, controlo de permissoes e varios modulos de gestao relacionados com uma frota. A plataforma permite centralizar informacao e demonstrar competencias de programacao, bases de dados, seguranca e organizacao de interfaces.")

    add_heading(doc, "9.2 Dificuldades Encontradas", 2)
    add_para(doc, "As maiores dificuldades estiveram relacionadas com a definicao da base de dados, a organizacao dos modulos e a necessidade de manter regras de acesso coerentes. Tambem foi desafiante garantir que os formularios recolhiam a informacao necessaria sem tornar a utilizacao demasiado complexa.")

    add_heading(doc, "9.3 Melhorias Futuras", 2)
    add_bullets(doc, [
        "Adicionar graficos estatisticos no dashboard.",
        "Melhorar a exportacao de relatorios em PDF.",
        "Criar notificacoes para manutencoes, cartas de conducao e prazos importantes.",
        "Adicionar mapas com localizacao de clientes ou viaturas.",
        "Reforcar validacoes e mensagens de erro.",
        "Melhorar a experiencia em dispositivos moveis.",
        "Preparar alojamento num servidor remoto.",
    ])

    add_heading(doc, "9.4 Conclusao Pessoal", 2)
    add_para(doc, "A realizacao do Frotalink permitiu consolidar conhecimentos adquiridos ao longo do curso, especialmente nas areas de desenvolvimento web, bases de dados e organizacao de sistemas de informacao. O projeto mostrou a importancia de planear bem a estrutura de dados, pensar na experiencia do utilizador e garantir seguranca nas operacoes principais.")
    add_para(doc, "Apesar das dificuldades, o resultado final apresenta uma base funcional e coerente para uma aplicacao de gestao de frotas. O Frotalink demonstra que uma solucao simples pode ajudar a resolver problemas reais de organizacao, consulta e acompanhamento de informacao operacional.")

    add_heading(doc, "10. Bibliografia", 1)
    add_para(doc, "PHP Manual. Disponivel em: https://www.php.net/manual/")
    add_para(doc, "MySQL Documentation. Disponivel em: https://dev.mysql.com/doc/")
    add_para(doc, "MDN Web Docs - HTML, CSS e JavaScript. Disponivel em: https://developer.mozilla.org/")
    add_para(doc, "XAMPP Apache Friends. Disponivel em: https://www.apachefriends.org/")
    add_para(doc, "OWASP - Cross-Site Request Forgery Prevention Cheat Sheet. Disponivel em: https://owasp.org/")

    add_heading(doc, "11. Anexos", 1)
    add_heading(doc, "11.1 Lista de Ficheiros Principais", 2)
    add_para(doc, "Este anexo pode ser utilizado para inserir uma captura da pasta do projeto ou uma listagem dos ficheiros principais.")
    add_heading(doc, "11.2 Capturas de Ecra", 2)
    add_para(doc, "Este anexo pode incluir capturas adicionais da aplicacao em funcionamento.")
    add_heading(doc, "11.3 Exportacao da Base de Dados", 2)
    add_para(doc, "Este anexo pode incluir o ficheiro SQL exportado atraves do phpMyAdmin ou uma descricao das tabelas criadas.")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    configure_styles(doc)
    add_footer(section)
    add_front_matter(doc)
    add_main_content(doc)
    apply_portuguese_accents(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
